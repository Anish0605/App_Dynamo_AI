"""
Generate Dynamo AI Architecture Document — Word + Diagram
"""

import matplotlib
matplotlib.use("Agg")
import matplotlib.pyplot as plt
import matplotlib.patches as mpatches
from matplotlib.patches import FancyBboxPatch, FancyArrowPatch
import numpy as np
from docx import Document
from docx.shared import Inches, Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import os

# ─────────────────────────────────────────────────────────────────────────────
# 1. ARCHITECTURAL DIAGRAM
# ─────────────────────────────────────────────────────────────────────────────

def make_diagram():
    W, H = 28, 20
    fig, ax = plt.subplots(1, 1, figsize=(W, H))
    ax.set_xlim(0, W)
    ax.set_ylim(0, H)
    ax.axis("off")
    fig.patch.set_facecolor("#0d1117")
    ax.set_facecolor("#0d1117")

    def box(x, y, w, h, label, sublabel="", color="#1e293b",
            text_color="white", fontsize=11, border="#334155", sublabel_size=9):
        rect = FancyBboxPatch((x, y), w, h,
                              boxstyle="round,pad=0.12",
                              facecolor=color, edgecolor=border, linewidth=2.0)
        ax.add_patch(rect)
        cy = y + h / 2 + (0.22 if sublabel else 0)
        ax.text(x + w/2, cy, label, ha="center", va="center",
                fontsize=fontsize, fontweight="bold", color=text_color,
                multialignment="center")
        if sublabel:
            ax.text(x + w/2, cy - 0.48, sublabel, ha="center", va="center",
                    fontsize=sublabel_size, color="#cbd5e1", multialignment="center")

    def arrow(x1, y1, x2, y2, color="#475569", lw=2.5):
        ax.annotate("", xy=(x2, y2), xytext=(x1, y1),
                    arrowprops=dict(arrowstyle="-|>", color=color, lw=lw,
                                   mutation_scale=18))

    # ── Title ──
    ax.text(W/2, 19.55, "Dynamo AI — Full System Architecture",
            ha="center", va="center", fontsize=20, fontweight="bold", color="white")
    ax.text(W/2, 19.1, "Academic Research Operating System  ·  FastAPI + Gemini + Supabase + Firebase + APIMart",
            ha="center", va="center", fontsize=12, color="#94a3b8")

    # ── Layer band definitions ──
    # (x, y, w, h, fill, edge, label, label_color)
    bands = [
        (0.4, 16.3, W-0.8, 2.5,  "#0c1a2e", "#2563eb", "LAYER 1 — FRONTEND",           "#60a5fa"),
        (0.4, 11.8, W-0.8, 4.2,  "#071d14", "#059669", "LAYER 2 — API GATEWAY",         "#34d399"),
        (0.4,  7.0, W-0.8, 4.5,  "#140d24", "#7c3aed", "LAYER 3 — AI MODELS",           "#c084fc"),
        (0.4,  3.5, W-0.8, 3.2,  "#1c1006", "#d97706", "LAYER 4 — EXTERNAL SERVICES",   "#fbbf24"),
        (0.4,  0.4, W-0.8, 2.8,  "#1a0510", "#e11d48", "LAYER 5 — QUOTA ENGINE",        "#fb7185"),
    ]
    for bx, by, bw, bh, fc, ec, bl, blc in bands:
        rect = FancyBboxPatch((bx, by), bw, bh, boxstyle="round,pad=0.12",
                              facecolor=fc, edgecolor=ec, linewidth=2.0, alpha=0.55)
        ax.add_patch(rect)
        ax.text(bx + 0.55, by + bh/2, bl, fontsize=10, color=blc,
                fontweight="bold", rotation=90, ha="center", va="center")

    # ─── LAYER 1: Frontend ───────────────────────────────────────────────────
    fe_y = 16.65
    fe_h = 1.7
    fe_gap = 0.25
    fe_x0  = 1.4
    fe_w   = (W - fe_x0 - 0.6 - 5*fe_gap) / 6

    fe_items = [
        ("Chat Interface",  "Markdown · LaTeX\nCode blocks"),
        ("Study Tools",     "Flashcards · Quiz\nRadio Mode"),
        ("Research Hub",    "Deep Research\nWatcher · Citations"),
        ("Create Tools",    "Images · Video\nDeck · Mindmap"),
        ("AI Detector",     "AI Score · Plagiarism\nHumanizer"),
        ("Profile / Auth",  "Firebase Auth\nQuota · Billing"),
    ]
    for i, (lbl, sub) in enumerate(fe_items):
        bx = fe_x0 + i*(fe_w + fe_gap)
        box(bx, fe_y, fe_w, fe_h, lbl, sub,
            color="#1e3a5f", text_color="white", fontsize=12, border="#3b82f6", sublabel_size=9.5)

    ax.text(W/2, 16.45,
            "Static HTML / Tailwind CSS / Vanilla JS  ·  Served on port 5000 via FastAPI StaticFiles",
            ha="center", va="center", fontsize=10, color="#60a5fa")

    # arrow down
    arrow(W/2, 16.3, W/2, 15.95)

    # ─── LAYER 2: API Gateway ────────────────────────────────────────────────
    # Main FastAPI box
    box(1.4, 14.45, 12.0, 1.35,
        "FastAPI Backend  (main.py — single process, port 5000)",
        "Auth check  ›  Quota gate  ›  Plan gate  ›  Keyword router  ›  Mode selector  ›  AI call  ›  Memory extract  ›  Save to DB",
        color="#064e3b", text_color="white", fontsize=13, border="#10b981", sublabel_size=10)

    # Mode routing boxes (right side)
    modes = [
        ("Fast\nMode",     "Free: lite model\nPlus/Pro: 3.5-flash"),
        ("DeepThink",      "Pro only\ndeep reasoning"),
        ("Research\nMode", "Plus/Pro\nClaude→Gemini→GPT"),
        ("Deep\nResearch", "Pro only\nAgentic agent"),
    ]
    m_x0  = 13.8
    m_w   = (W - m_x0 - 0.6 - 3*0.22) / 4
    m_y   = 13.0
    m_h   = 2.75
    for i, (lbl, sub) in enumerate(modes):
        bx = m_x0 + i*(m_w + 0.22)
        box(bx, m_y, m_w, m_h, lbl, sub,
            color="#065f46", text_color="#86efac", fontsize=11, border="#10b981", sublabel_size=9)

    # Dispatcher bar spanning full width
    box(1.4, 12.0, W-2.0, 0.85,
        "Mode Dispatcher  ·  Plan Enforcement Gates  ·  Quota Enforcer  ·  Keyword Router",
        color="#022c22", text_color="#6ee7b7", fontsize=12, border="#059669")

    arrow(W/2, 12.0, W/2, 11.45)

    # ─── LAYER 3: AI Models ──────────────────────────────────────────────────
    model_h = 2.0
    model_y = 8.9

    # Sub-label: API group bars
    box(1.4, 7.15, 9.2, 0.65, "Google GenAI SDK  (Direct API)",
        color="#1e1035", text_color="#c084fc", fontsize=11, border="#7c3aed")
    box(10.8, 7.15, 4.1, 0.65, "Gemini Interactions API",
        color="#1e1035", text_color="#c084fc", fontsize=11, border="#7c3aed")
    box(15.1, 7.15, W-15.7, 0.65, "APIMart Gateway  (with Gemini 3.5-flash fallback)",
        color="#1e1035", text_color="#c084fc", fontsize=11, border="#7c3aed")

    models = [
        (1.4,  "Gemini 3.1\nFlash Lite",          "Fast Mode (Free)\nUtility tasks\nWatcher · Deck",   "#2d1260"),
        (5.6,  "Gemini 3.5\nFlash",               "Fast (Plus/Pro)\nDeepThink · Detector\nFallback",    "#2d1260"),
        (9.8,  "Gemini\nDeep Research",            "deep-research-\npreview-04-2026\nPro only",          "#2d1260"),
        (13.5, "Claude\nSonnet 4.5",               "Research Step 1\nExtract & structure",               "#1a2850"),
        (17.2, "Gemini 3.1\n(APIMart)",            "Research Step 2\nAnalyse & synthesise",              "#1a2850"),
        (21.0, "GPT-5.4\n(APIMart)",               "Research Step 3\nWrite final report",                "#1a2850"),
    ]
    m_fw = 3.7
    for mx, lbl, sub, c in models:
        box(mx, model_y, m_fw, model_h, lbl, sub,
            color=c, text_color="#e9d5ff", fontsize=12, border="#a855f7", sublabel_size=9.5)

    arrow(W/2, 8.9, W/2, 8.3)

    # Specialised modules row
    spec_items = [
        ("Memory\nEngine",     "Extracts & injects\npersonal context\n(Plus/Pro)"),
        ("Document\nLibrary",  "PDF/DOCX/TXT\nsummarise & inject\ninto prompt"),
        ("Plagiarism\nChecker","Tavily + Semantic\nScholar + Gemini\nscoring"),
        ("Citation\nChecker",  "8 format styles\nDOI verify\nAuto-correct"),
        ("Research\nWatcher",  "Daily monitor\nBrevo email digest\n(Pro only)"),
        ("Content\nFactory",   "Mindmap · Flowchart\nDeck · Flashcards\nQuiz"),
        ("Media\nEngine",      "Image (Pollinations)\nVideo · Voice\nEdge TTS"),
    ]
    s_x0 = 1.4
    s_w  = (W - s_x0 - 0.6 - 6*0.22) / 7
    s_y  = 7.17
    s_h  = 1.55
    for i, (lbl, sub) in enumerate(spec_items):
        bx = s_x0 + i*(s_w + 0.22)
        box(bx, s_y, s_w, s_h, lbl, sub,
            color="#1c1917", text_color="#d1d5db", fontsize=11, border="#52525b", sublabel_size=9)

    arrow(W/2, 7.17, W/2, 6.65)

    # ─── LAYER 4: External Services ─────────────────────────────────────────
    svcs = [
        ("Supabase\nPostgreSQL", "users · memories\ndocuments · chats\nsubscriptions"),
        ("Firebase\nAuth",       "Google OAuth\nEmail+Password\nUID bridge"),
        ("Tavily\nSearch",       "Live web results\nDeep search\n(Research pipeline)"),
        ("Semantic\nScholar",    "200M+ papers\nAcademic search\nDOI resolution"),
        ("Razorpay\nPayments",   "Plus ₹399/mo\nPro ₹999/mo\nHMAC verify"),
        ("Brevo\nEmail",         "Watcher digests\nTransactional\nSMTP"),
        ("Pollinations\nAI",     "Image generation\nFree API\nPlus/Pro only"),
    ]
    s2_x0 = 1.4
    s2_w  = (W - s2_x0 - 0.6 - 6*0.22) / 7
    s2_y  = 3.75
    s2_h  = 2.65
    for i, (lbl, sub) in enumerate(svcs):
        bx = s2_x0 + i*(s2_w + 0.22)
        box(bx, s2_y, s2_w, s2_h, lbl, sub,
            color="#1c0a00", text_color="#fde68a", fontsize=12, border="#d97706", sublabel_size=9.5)

    arrow(W/2, 3.75, W/2, 3.2)

    # ─── LAYER 5: Quota Engine ───────────────────────────────────────────────
    box(1.4, 0.65, W-2.0, 2.35,
        "Freemium Quota Engine  (supabase_client.py)",
        "Free: 10 msg/day  ·  Plus: 100 msg/day (₹399/mo)  ·  Pro: 300 msg/day (₹999/mo)\n"
        "Image caps: 0 / 25 / 100 per month  ·  Video caps: 0 / 5 / 25 per month  ·  Lazy daily reset at midnight UTC",
        color="#3b0614", text_color="#fda4af", fontsize=13, border="#e11d48", sublabel_size=10.5)

    # ── Legend ──────────────────────────────────────────────────────────────
    legend = [
        ("#2563eb", "Frontend"),
        ("#059669", "FastAPI"),
        ("#7c3aed", "AI Models"),
        ("#52525b", "Modules"),
        ("#d97706", "Services"),
        ("#e11d48", "Quota"),
    ]
    lx0, ly = 1.4, 0.27
    for i, (c, lbl) in enumerate(legend):
        bx = lx0 + i * 4.3
        ax.add_patch(mpatches.Rectangle((bx, ly), 0.45, 0.22, color=c))
        ax.text(bx + 0.58, ly + 0.11, lbl, fontsize=10, color="white", va="center")

    plt.tight_layout(pad=0.2)
    path = "/home/runner/workspace/architecture_diagram.png"
    plt.savefig(path, dpi=200, bbox_inches="tight", facecolor=fig.get_facecolor())
    plt.close()
    print(f"Diagram saved: {path}")
    return path

# ─────────────────────────────────────────────────────────────────────────────
# 2. WORD DOCUMENT
# ─────────────────────────────────────────────────────────────────────────────

def add_heading(doc, text, level=1, color=None):
    h = doc.add_heading(text, level=level)
    h.alignment = WD_ALIGN_PARAGRAPH.LEFT
    if color:
        for run in h.runs:
            run.font.color.rgb = RGBColor(*color)
    return h

def add_para(doc, text, bold=False, italic=False, size=10, color=None, align=None):
    p = doc.add_paragraph()
    if align:
        p.alignment = align
    run = p.add_run(text)
    run.bold = bold
    run.italic = italic
    run.font.size = Pt(size)
    if color:
        run.font.color.rgb = RGBColor(*color)
    return p

def add_table_row(table, cells, bold_first=False):
    row = table.add_row()
    for i, (cell, text) in enumerate(zip(row.cells, cells)):
        cell.text = text
        if bold_first and i == 0:
            for para in cell.paragraphs:
                for run in para.runs:
                    run.bold = True
    return row

def set_col_widths(table, widths):
    for row in table.rows:
        for i, cell in enumerate(row.cells):
            if i < len(widths):
                cell.width = widths[i]

def build_doc(diagram_path):
    doc = Document()

    # Page margins
    for section in doc.sections:
        section.top_margin    = Cm(2.0)
        section.bottom_margin = Cm(2.0)
        section.left_margin   = Cm(2.5)
        section.right_margin  = Cm(2.5)

    # ── Title page ──
    doc.add_paragraph()
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = title.add_run("Dynamo AI: Architecture and Design of an\nAI-Powered Academic Research Operating System")
    r.bold = True
    r.font.size = Pt(20)
    r.font.color.rgb = RGBColor(15, 23, 42)

    doc.add_paragraph()
    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    rs = sub.add_run("Technical Architecture Paper  |  May 2026")
    rs.italic = True
    rs.font.size = Pt(12)
    rs.font.color.rgb = RGBColor(100, 116, 139)

    doc.add_paragraph()

    # ── Abstract ──
    add_heading(doc, "Abstract", level=1)
    add_para(doc,
        "Dynamo AI is a full-stack, AI-powered Academic Research Operating System designed for students, "
        "researchers, and professionals in India. Built on a FastAPI Python backend and a static HTML/JS frontend, "
        "it integrates multiple frontier AI models — including Google Gemini 3.5 Flash, Claude Sonnet 4.5, and "
        "GPT-5.4 — behind a unified chat interface. The platform delivers a differentiated freemium experience "
        "across three subscription tiers, enforcing model-level differentiation, plan-gated features, and per-user "
        "daily/monthly quotas. This paper describes the full system architecture, each subsystem's design and "
        "data flow, the AI model routing strategy, cost-optimisation decisions, and the external service integrations "
        "that power Dynamo AI's research, detection, memory, and content-creation capabilities.",
        size=10)

    doc.add_paragraph()

    # ── 1. Introduction ──
    add_heading(doc, "1. Introduction", level=1)
    add_para(doc,
        "The rapid proliferation of large language models (LLMs) has created new opportunities for building "
        "intelligent academic tools. Dynamo AI was conceived as a Research Operating System — not merely a "
        "chatbot, but a platform that combines conversational AI with web search, academic databases, document "
        "memory, plagiarism detection, citation verification, and structured research pipelines, all behind a "
        "single interface optimised for Indian academic users.", size=10)
    add_para(doc,
        "The system is intentionally lean: a single FastAPI process serves both the static frontend and all "
        "API endpoints on port 5000, keeping operational complexity low while supporting a rich feature set. "
        "Model selection, quota enforcement, and feature gating are all handled server-side, ensuring that "
        "the business logic cannot be bypassed by frontend manipulation.", size=10)

    # ── 2. High-Level Architecture ──
    add_heading(doc, "2. High-Level Architecture", level=1)
    add_para(doc,
        "Dynamo AI follows a five-layer architecture as illustrated in Figure 1:", size=10)

    layers_desc = [
        ("Layer 1 — Frontend", "Static HTML pages with Tailwind CSS and vanilla JavaScript modules. "
         "No build step or framework. Pages are served directly by FastAPI's StaticFiles middleware."),
        ("Layer 2 — API Gateway (FastAPI)", "A single FastAPI application (main.py) handles all HTTP "
         "requests, enforces authentication, checks quotas, applies plan gates, routes to the correct "
         "AI mode, and assembles the final response."),
        ("Layer 3 — AI Model Tier", "Multiple Gemini models via the Google GenAI SDK, the Gemini "
         "Interactions API for deep research, and Claude/GPT-5.4 accessed through the APIMart gateway."),
        ("Layer 4 — External Services", "Tavily (web search), Semantic Scholar (academic papers), "
         "Supabase (PostgreSQL database), Firebase (authentication), Razorpay (payments), Brevo (email), "
         "Pollinations AI (image generation), and Microsoft Edge TTS (text-to-speech)."),
        ("Layer 5 — Quota Engine", "A plan-aware quota system backed by Supabase that enforces daily "
         "message limits, monthly image/video caps, and per-feature plan gates."),
    ]

    for title_t, desc in layers_desc:
        p = doc.add_paragraph(style="List Bullet")
        run_b = p.add_run(title_t + ": ")
        run_b.bold = True
        run_b.font.size = Pt(10)
        run_n = p.add_run(desc)
        run_n.font.size = Pt(10)

    doc.add_paragraph()

    # Insert diagram
    add_para(doc, "Figure 1: Dynamo AI Full System Architecture Diagram", bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, size=9)
    pic_para = doc.add_paragraph()
    pic_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_pic = pic_para.add_run()
    run_pic.add_picture(diagram_path, width=Inches(6.5))

    doc.add_paragraph()

    # ── 3. Frontend Architecture ──
    add_heading(doc, "3. Frontend Architecture", level=1)
    add_para(doc,
        "The frontend is a collection of static files — primarily a single large Index.html with modular "
        "JavaScript files loaded via cache-busted <script> tags. Tailwind CSS (CDN in development, "
        "intended for PostCSS in production) provides the styling. No React, Vue, or build toolchain "
        "is required, which means the frontend loads instantly with zero JavaScript bundle compilation.", size=10)

    add_heading(doc, "3.1 Key JavaScript Modules", level=2)
    modules_fe = [
        ("chat.js", "Core chat loop: sends messages, streams/renders responses, handles smart actions (summarise, explain), quota checks client-side."),
        ("ui.js", "Composer bar, mode switching (Fast/DeepThink/Research), tools dropdown, flyout menus, hero animation."),
        ("sidebar.js", "Chat history, folder management, tabbed sidebar (Chats / Folders), search/filter."),
        ("detector.js", "AI Detector modal (AI score, sentence heatmap, humanizer) and Plagiarism modal (originality check, self-plagiarism). Passes user_id for backend plan gating."),
        ("deep_research.js", "In-chat Deep Research Agent interface: starts jobs, polls for progress, renders the final report."),
        ("citation_checker.js", "Citation analysis modal with 8 format styles, live DOI verification, and auto-correction export."),
        ("profile.js", "Profile modal: plan display, live credit pills (messages/images/videos used today vs limit), name/email/password editing."),
        ("memory.js", "AI Memory viewer — lists extracted memories, allows deletion."),
        ("documents.js", "Document Library modal — upload, list, delete saved documents."),
        ("write_paper.js", "Research paper writer — topic input, citation format selector, multi-model pipeline trigger."),
    ]

    tbl = doc.add_table(rows=1, cols=2)
    tbl.style = "Table Grid"
    hdr = tbl.rows[0].cells
    hdr[0].text = "Module"
    hdr[1].text = "Responsibility"
    for cell in hdr:
        for para in cell.paragraphs:
            for run in para.runs:
                run.bold = True
    for mod, desc in modules_fe:
        row = tbl.add_row()
        row.cells[0].text = mod
        row.cells[1].text = desc
    set_col_widths(tbl, [Inches(1.5), Inches(5.0)])
    doc.add_paragraph()

    add_heading(doc, "3.2 Example — Fast Mode Chat Flow (Frontend Side)", level=2)
    add_para(doc,
        "When a user types a message and presses send in Fast Mode, the following sequence occurs:", size=10)
    steps_fe = [
        "chat.js calls checkMessageLimit() — compares localStorage counter against the plan limit (10/100/300). If over limit, shows an upgrade card immediately without hitting the server.",
        "chat.js calls window.callBackend('/chat', payload) with: { message, chat_id, user_id, mode:'chat', deep_dive:false, use_search:false }.",
        "The response JSON is received and renderMarkdown() converts it to formatted HTML with syntax-highlighted code blocks, LaTeX support, and source citations.",
        "The sidebar is refreshed via loadChatSidebar() and the daily counter is incremented in localStorage.",
    ]
    for i, step in enumerate(steps_fe, 1):
        p = doc.add_paragraph(style="List Number")
        p.add_run(step).font.size = Pt(10)
    doc.add_paragraph()

    # ── 4. Backend Architecture ──
    add_heading(doc, "4. Backend Architecture (FastAPI)", level=1)
    add_para(doc,
        "The backend is a single FastAPI application (backend/main.py) with approximately 1,500 lines "
        "that handles every API endpoint. It is launched by Uvicorn and serves both the static frontend "
        "and all REST endpoints on port 5000.", size=10)

    add_heading(doc, "4.1 Request Processing Pipeline", level=2)
    pipeline_steps = [
        ("Step 1: Authentication", "The client sends its Supabase user_id in the request body. The backend calls get_user_by_supabase_id() which fetches the full user record from Supabase (including plan, daily_quota_used, quota_date). This also triggers a daily quota reset if the stored quota_date differs from today (UTC)."),
        ("Step 2: Quota Enforcement", "check_user_quota(user) compares daily_quota_used against PLAN_LIMITS[plan]['daily_chat'] (Free:10, Plus:100, Pro:300). If exceeded, an error response is returned immediately before any AI call is made."),
        ("Step 3: Plan Gate", "For gated features (Research Mode, DeepThink, Deep Research Agent, AI Detector, Memory), the user's plan is checked against the minimum required plan. Free users receive an HTTP 403 or a structured error response with an upgrade link."),
        ("Step 4: Keyword Routing", "The message is scanned for keywords: image generation triggers (\"generate image\"), video triggers (\"create video\"), flowchart triggers (\"flowchart\", \"steps\", \"workflow\"), mindmap triggers. smart_action:true bypasses this routing."),
        ("Step 5: Web Search (conditional)", "If use_search:true or the mode is Research, Tavily fetches live web context (3–5 sources for Fast, 10+ for Research) that is appended to the system prompt."),
        ("Step 6: AI Response", "model.get_ai_response() selects the model based on plan and deep_dive flag, builds the full prompt (system + memory context + document library + search context + conversation history), and calls the Gemini API."),
        ("Step 7: Memory Extraction", "After the response, for Plus/Pro users only, _should_extract_memory() runs a two-layer filter (social blocklist + personal signal regex). If the message contains genuine personal context, memory.extract_memories() is called in a background thread."),
        ("Step 8: Quota Increment", "increment_quota(user) writes daily_quota_used + 1 back to Supabase."),
    ]
    for title_s, desc in pipeline_steps:
        p = doc.add_paragraph(style="List Bullet")
        run_b = p.add_run(title_s + ": ")
        run_b.bold = True
        run_b.font.size = Pt(10)
        p.add_run(desc).font.size = Pt(10)
    doc.add_paragraph()

    # ── 5. AI Model Routing Strategy ──
    add_heading(doc, "5. AI Model Routing Strategy", level=1)
    add_para(doc,
        "One of Dynamo AI's defining architectural decisions is its tiered, plan-aware model routing. "
        "Rather than serving the same AI model to all users, the backend dynamically selects the "
        "appropriate model based on the request type and the user's subscription plan:", size=10)

    doc.add_paragraph()
    tbl2 = doc.add_table(rows=1, cols=4)
    tbl2.style = "Table Grid"
    hdr2 = tbl2.rows[0].cells
    for cell, text in zip(hdr2, ["Mode", "Model", "Who Gets It", "Rationale"]):
        cell.text = text
        for para in cell.paragraphs:
            for run in para.runs:
                run.bold = True
    model_rows = [
        ("Fast Mode", "gemini-3.1-flash-lite-preview", "Free users", "Cost-optimised; handles the highest message volume."),
        ("Fast Mode", "gemini-3.5-flash", "Plus & Pro users", "Frontier quality for paying users, even in quick mode."),
        ("DeepThink Mode", "gemini-3.5-flash + deep reasoning prompt", "Pro only", "Same model as Fast (Plus/Pro), but with a structured system prompt that forces step-by-step reasoning, multiple perspectives, and gap analysis."),
        ("Research Mode", "APIMart pipeline: Claude Sonnet 4.5 → Gemini 3.1 → GPT-5.4 (Gemini 3.5-flash fallback)", "Plus & Pro", "Multi-model consensus for long-form research. Each model handles a specialised role: Claude extracts, Gemini analyses, GPT writes."),
        ("Deep Research Agent", "deep-research-preview-04-2026 / deep-research-max-preview-04-2026", "Pro only", "Gemini Interactions API — performs true agentic research with web browsing, multi-step reasoning, and gap analysis."),
        ("Utility tasks", "gemini-3.1-flash-lite-preview", "All plans", "Flashcards, mindmaps, flowcharts, voice transcription, document summarisation — quality-insensitive background tasks."),
        ("AI Detector / Plagiarism", "gemini-3.5-flash", "Plus & Pro", "Quality-critical: borderline AI scores and subtle paraphrase overlap require frontier reasoning."),
        ("Citation Checker", "gemini-3.5-flash → lite fallback", "Plus & Pro", "3.5-flash primary; lite fallback ensures resilience on quota exhaustion."),
        ("Research Watcher", "gemini-3.1-flash-lite-preview", "Pro only", "Runs once per day per user; lite model sufficient for monitoring summaries."),
    ]
    for row_data in model_rows:
        row = tbl2.add_row()
        for cell, text in zip(row.cells, row_data):
            cell.text = text
    doc.add_paragraph()

    add_heading(doc, "5.1 Example — DeepThink vs Fast Mode (Same User, Same Query)", level=2)
    add_para(doc, "Query: \"Explain the implications of transformer attention mechanisms for long-context reasoning.\"", bold=True, size=10)
    add_para(doc,
        "Fast Mode (Plus user): The system prompt is the standard Dynamo AI assistant prompt. "
        "Gemini 3.5-flash responds in ~2 seconds with a clear, well-structured explanation. "
        "No explicit reasoning chain is required.", size=10)
    add_para(doc,
        "DeepThink Mode (Pro user): The system prompt includes structured reasoning directives: "
        "\"Think step by step. Consider multiple angles. Identify what's missing. Challenge your own "
        "assumptions.\" The same Gemini 3.5-flash model produces a significantly deeper response — "
        "not because the model changed, but because the system prompt changed the output behaviour. "
        "The token count is higher, the latency is ~1 second more, but the analytical depth is substantially greater.", size=10)
    doc.add_paragraph()

    # ── 6. Research Mode Pipeline ──
    add_heading(doc, "6. Research Mode Multi-Model Pipeline", level=1)
    add_para(doc,
        "Research Mode (Plus/Pro) routes through a three-stage pipeline implemented in "
        "multi_model_router.py. The pipeline is designed for long-form, cited research reports "
        "rather than conversational answers:", size=10)

    pipeline_research = [
        ("Stage 1 — Web Search", "Tavily performs deep search (10+ results) on the user's topic. Results are compiled into a structured web_context block with titles, URLs, and content snippets."),
        ("Stage 2 — Extraction (Claude Sonnet 4.5 via APIMart)", "Claude receives the web context and the user's topic. It extracts the most relevant factual claims, identifies key arguments, and structures them into a research outline. Claude's strength in careful extraction and nuanced reading makes it ideal for this stage."),
        ("Stage 3 — Analysis (Gemini 3.1 via APIMart)", "Gemini receives Claude's extracted outline and performs analytical synthesis — comparing viewpoints, identifying consensus and controversy, and adding academic context."),
        ("Stage 4 — Writing (GPT-5.4 via APIMart)", "GPT-5.4 takes the synthesis and writes the final research report in the user's chosen citation format (APA 7th, MLA 9th, Chicago 17th, Harvard, IEEE, Vancouver, ACS, or ASA). It applies proper in-text citations and generates a formatted bibliography."),
        ("Fallback", "If any APIMart stage fails (timeout, rate limit, or API error), that stage falls back to Gemini 3.5-flash via the direct Gemini API. This ensures the pipeline never fully fails."),
    ]
    for title_s, desc in pipeline_research:
        p = doc.add_paragraph(style="List Bullet")
        p.add_run(title_s + ": ").bold = True
        p.add_run(desc).font.size = Pt(10)
    doc.add_paragraph()

    # ── 7. Memory System ──
    add_heading(doc, "7. AI Memory System", level=1)
    add_para(doc,
        "The Memory System (backend/memory.py) gives Dynamo AI persistent awareness of each user "
        "across sessions. It is available to Plus and Pro users only.", size=10)

    add_heading(doc, "7.1 Memory Extraction Filter", level=2)
    add_para(doc,
        "A two-layer filter runs after every AI response to decide whether memory extraction is warranted:", size=10)
    filter_steps = [
        "Layer 1 — Social Blocklist: If the message is exactly a social phrase (\"hi\", \"hello\", \"thanks\", \"ok\", \"great\", \"bye\", etc.) it is rejected immediately. This prevents wasted API calls on purely conversational filler.",
        "Layer 2 — Personal Signal Regex: The message is scanned for first-person academic/personal signals: \"I failed\", \"my exam\", \"I'm preparing\", \"I struggle with\", \"my university\", \"I got\", \"I need to learn\", etc. Only messages that contain genuine personal context trigger memory extraction.",
    ]
    for step in filter_steps:
        p = doc.add_paragraph(style="List Bullet")
        p.add_run(step).font.size = Pt(10)

    add_heading(doc, "7.2 Memory Storage and Injection", level=2)
    add_para(doc,
        "Extracted memories are stored as structured records in Supabase's user_memories table "
        "(user_id, memory_text, created_at). On every subsequent chat request, the user's memories "
        "are fetched and injected into the system prompt as a dedicated context block: "
        "\"--- What you know about this user --- ... --- End of user context ---\". "
        "The AI is instructed to use this context naturally without announcing it.", size=10)

    add_heading(doc, "7.3 Example", level=2)
    add_para(doc, "User (Session 1): \"I keep failing organic chemistry. My final exam is in 3 weeks.\"", italic=True, size=10)
    add_para(doc,
        "Memory extracted: \"User is struggling with organic chemistry and has a final exam in approximately 3 weeks.\"", size=10)
    add_para(doc, "User (Session 2, one week later): \"Explain the SN2 reaction mechanism.\"", italic=True, size=10)
    add_para(doc,
        "Dynamo AI responds with: \"Since you have your organic chemistry final coming up soon, let me walk you through SN2 in a way that will help you answer exam questions...\" — without the user having to re-explain their context.", size=10)
    doc.add_paragraph()

    # ── 8. Document Library ──
    add_heading(doc, "8. Document Library", level=1)
    add_para(doc,
        "The Document Library (backend/documents.py) allows users to upload PDFs, DOCX, or TXT files "
        "that Dynamo AI then summarises and permanently references in every future chat session. "
        "This is distinct from the file-upload-per-chat feature — the library is persistent.", size=10)

    add_para(doc,
        "On upload: The file text is extracted (via pdfplumber for PDFs, python-docx for DOCX), "
        "truncated to 8,000 characters, and passed to Gemini 3.1-flash-lite-preview with a prompt "
        "requesting a structured JSON summary containing: a 2–4 sentence plain-English summary, "
        "up to 8 key technical terms, and up to 4 subject areas.", size=10)

    add_para(doc,
        "On every chat: format_docs_for_prompt() builds a \"User's Saved Document Library\" block "
        "that is appended to the system prompt. The AI naturally references the user's documents "
        "when answering relevant questions without being explicitly told to.", size=10)
    doc.add_paragraph()

    # ── 9. AI Detector + Plagiarism ──
    add_heading(doc, "9. AI Detector and Plagiarism Checker", level=1)
    add_para(doc,
        "The AI Detector and Plagiarism Checker (backend/detector.py) is an in-house feature "
        "requiring no third-party detection API. It uses Gemini 3.5-flash, Tavily, and Semantic Scholar. "
        "It is available to Plus and Pro users only (backend-enforced).", size=10)

    detector_features = [
        ("AI Text Detection", "detect_ai()", "Gemini 3.5-flash analyses up to 4,000 characters of text for AI-generation patterns: burstiness, perplexity uniformity, filler phrases, lack of personal voice, and over-structured prose. Returns a score 0–100, a label (Human / Uncertain / AI-Generated), and 3–5 specific signals."),
        ("Sentence Heatmap", "detect_ai_sentences()", "Gemini 3.5-flash scores every sentence individually (0=human, 100=AI). The frontend renders a colour-coded heatmap over the text. Academic writing conventions (passive voice, hedging, citations) are explicitly excluded from penalisation."),
        ("Humanizer", "humanize_text()", "Gemini 3.5-flash rewrites AI-generated text to preserve the original meaning, academic register, and all citations, while introducing natural hedging, varied sentence structure, and personal voice markers."),
        ("Originality Check", "check_plagiarism()", "Three 40-word query phrases are extracted from the beginning, middle, and end of the submitted text. Tavily performs web searches; Semantic Scholar is queried for academic papers. All unique sources are deduplicated and passed to Gemini 3.5-flash, which scores the actual similarity (0–100) while explicitly distinguishing common terminology from plagiarised content."),
        ("Self-Plagiarism Check", "check_self_plagiarism()", "Two documents are submitted (current paper + prior work). Gemini 3.5-flash compares them directly (no external search), identifies specific overlapping passages, distinguishes boilerplate methodology from recycled content, and provides a score + specific overlaps list + practical recommendation."),
    ]

    tbl3 = doc.add_table(rows=1, cols=3)
    tbl3.style = "Table Grid"
    for cell, text in zip(tbl3.rows[0].cells, ["Feature", "Function", "How It Works"]):
        cell.text = text
        for para in cell.paragraphs:
            for run in para.runs:
                run.bold = True
    for feat, func, desc in detector_features:
        row = tbl3.add_row()
        row.cells[0].text = feat
        row.cells[1].text = func
        row.cells[2].text = desc
    doc.add_paragraph()

    # ── 10. Deep Research Agent ──
    add_heading(doc, "10. Deep Research Agent", level=1)
    add_para(doc,
        "The Deep Research Agent (backend/deep_research.py) is Dynamo AI's most powerful research "
        "capability, available to Pro users only. It uses Google's Gemini Interactions API — a "
        "dedicated endpoint for agentic, multi-step web research.", size=10)

    add_para(doc,
        "Two model variants are available: deep-research-preview-04-2026 (standard) and "
        "deep-research-max-preview-04-2026 (extended, deeper research). The agent autonomously "
        "browses the web, evaluates sources, synthesises findings, and produces a structured "
        "research report — similar to a junior research assistant completing a literature review.", size=10)

    add_heading(doc, "10.1 Job Architecture", level=2)
    add_para(doc,
        "Research jobs are asynchronous: the frontend starts a job and receives a job_id. "
        "It then polls GET /deep-research-status/{job_id} every 3 seconds to check progress. "
        "Jobs are stored in an in-memory dictionary keyed by job_id, with fields: status, "
        "progress_msg, activity log, elapsed time, and the final report. A fallback pipeline "
        "(6 Tavily searches + gap analysis + Gemini synthesis) activates if the Interactions API "
        "is unavailable.", size=10)
    doc.add_paragraph()

    # ── 11. Research Watcher ──
    add_heading(doc, "11. Research Watcher", level=1)
    add_para(doc,
        "The Research Watcher (backend/watcher_check.py, watcher_scheduler.py) is a Pro-only "
        "background monitoring service that checks user-defined research topics once every 24 hours "
        "and sends email digests via Brevo when new developments are found.", size=10)

    watcher_steps = [
        "Users define topics to watch (e.g., \"AI safety research\", \"quantum computing breakthroughs\").",
        "Every 24 hours, the scheduler (APScheduler) triggers a check for each Pro user's active watches.",
        "For each watch: Tavily fetches recent web results on the topic. Gemini 3.1-flash-lite-preview evaluates whether the results represent genuinely new developments (not just re-indexed old content).",
        "If new content is found, Brevo's transactional email API sends a formatted digest to the user's email address with source links and a summary.",
    ]
    for step in watcher_steps:
        p = doc.add_paragraph(style="List Number")
        p.add_run(step).font.size = Pt(10)
    doc.add_paragraph()

    # ── 12. Freemium / Quota System ──
    add_heading(doc, "12. Freemium and Quota System", level=1)
    add_para(doc,
        "Dynamo AI's monetisation is built around a three-tier freemium model enforced at both "
        "the frontend (for UX) and backend (for security):", size=10)

    tbl4 = doc.add_table(rows=1, cols=5)
    tbl4.style = "Table Grid"
    for cell, text in zip(tbl4.rows[0].cells, ["Plan", "Daily Chat", "Images/Mo", "Videos/Mo", "Price"]):
        cell.text = text
        for para in cell.paragraphs:
            for run in para.runs:
                run.bold = True
    for row_data in [
        ("Free", "10 messages", "0", "0", "₹0 / forever"),
        ("Plus", "100 messages", "25", "5", "₹399 / month"),
        ("Pro", "300 messages", "100", "25", "₹999 / month"),
    ]:
        row = tbl4.add_row()
        for cell, text in zip(row.cells, row_data):
            cell.text = text
    doc.add_paragraph()

    add_heading(doc, "12.1 Quota Reset Logic", level=2)
    add_para(doc,
        "Daily quotas reset at midnight UTC (approximately 5:30 AM IST). The reset is lazy — "
        "it happens on the next API request after midnight, not on a scheduled cron. "
        "If quota_date in the user's Supabase record differs from today's UTC date, "
        "daily_quota_used is reset to 0 and quota_date is updated atomically.", size=10)

    add_heading(doc, "12.2 Plan Gates", level=2)
    gates = [
        ("Research Mode", "Plus and Pro"),
        ("AI Memory", "Plus and Pro"),
        ("Document Library", "Plus and Pro"),
        ("AI Detector / Plagiarism Checker", "Plus and Pro"),
        ("Image Generation", "Plus and Pro"),
        ("Video Generation", "Plus and Pro"),
        ("Citation Checker", "Plus and Pro"),
        ("DeepThink Mode", "Pro only"),
        ("Find Research Gaps", "Pro only"),
        ("Deep Research Agent", "Pro only"),
        ("Research Watcher", "Pro only"),
    ]
    tbl5 = doc.add_table(rows=1, cols=2)
    tbl5.style = "Table Grid"
    for cell, text in zip(tbl5.rows[0].cells, ["Feature", "Minimum Plan"]):
        cell.text = text
        for para in cell.paragraphs:
            for run in para.runs:
                run.bold = True
    for feat, plan in gates:
        row = tbl5.add_row()
        row.cells[0].text = feat
        row.cells[1].text = plan
    doc.add_paragraph()

    # ── 13. Payment Integration ──
    add_heading(doc, "13. Payment Integration (Razorpay)", level=1)
    add_para(doc,
        "Payments are handled via Razorpay (backend/payments.py). The flow is:", size=10)
    payment_steps = [
        "Frontend calls POST /create-order with the selected plan (plus/pro). Backend creates a Razorpay order and returns order_id and amount.",
        "Frontend opens the Razorpay checkout modal. On success, Razorpay returns razorpay_payment_id, razorpay_order_id, and razorpay_signature.",
        "Frontend calls POST /verify-payment. Backend verifies the HMAC-SHA256 signature using the Razorpay secret key. On success, it updates users.plan in Supabase and inserts a row into the subscriptions table.",
        "Razorpay also sends async webhook events to POST /webhook as a secondary safety net — the backend re-verifies and re-applies any plan updates that may have been missed.",
    ]
    for step in payment_steps:
        p = doc.add_paragraph(style="List Number")
        p.add_run(step).font.size = Pt(10)
    doc.add_paragraph()

    # ── 14. Cost Optimisation ──
    add_heading(doc, "14. Cost Optimisation Strategy", level=1)
    add_para(doc,
        "Following a ₹1,400 Gemini API bill in May 2026 (a +4,970% increase vs April), a systematic "
        "cost-reduction audit was conducted. The key decisions made:", size=10)

    cost_items = [
        ("Plan-aware model routing", "Free users get gemini-3.1-flash-lite-preview for Fast Mode; Plus/Pro get gemini-3.5-flash. The majority of message volume is from free users during onboarding — this one change provides the largest per-message cost reduction."),
        ("Memory extraction gating", "Memory extraction now runs only for Plus/Pro users, and only when the message passes a two-layer personal-context filter. Free users never trigger memory API calls."),
        ("Feature gating at backend", "All premium features are enforced server-side, not just in the frontend. A user cannot bypass the frontend and call expensive endpoints directly."),
        ("Watcher model downgrade", "Research Watcher switched from gemini-2.5-flash-preview to gemini-3.1-flash-lite-preview (monitoring summaries don't require frontier intelligence) and frequency reduced from hourly to once per 24 hours."),
        ("Utility task models", "Background and utility tasks (flashcards, mindmaps, flowcharts, deck generation, document summarisation, voice transcription, file analysis) all use gemini-3.1-flash-lite-preview."),
        ("Detector model upgrade (quality trade-off)", "The AI Detector was upgraded from gemini-3-flash-preview to gemini-3.5-flash — a deliberate cost increase justified by the accuracy requirement: users trust the plagiarism and AI detection scores with their academic work."),
    ]
    for title_c, desc in cost_items:
        p = doc.add_paragraph(style="List Bullet")
        p.add_run(title_c + ": ").bold = True
        p.add_run(desc).font.size = Pt(10)
    doc.add_paragraph()

    # ── 15. Data Architecture ──
    add_heading(doc, "15. Data Architecture (Supabase / PostgreSQL)", level=1)

    db_tables = [
        ("users", "Primary user record: firebase_uid, plan, full_name, email, daily_quota_used, quota_date, image_count_used, video_count_used, quota_month, paper_count_used."),
        ("chat_messages", "All chat messages keyed by chat_id, with role (user/assistant) and content."),
        ("user_memories", "Extracted memory records: user_id, memory_text, created_at. Used to personalise future sessions."),
        ("user_documents", "Document library: user_id, filename, summary, key_terms, topics, upload_date, file_size_kb."),
        ("subscriptions", "Payment records: user_id, plan, razorpay_order_id, razorpay_payment_id, amount, status, expires_at."),
        ("folders", "Chat folder organisation: user_id, name. Chats can be assigned to folders."),
        ("research_watches", "User-defined research topics for the Research Watcher: user_id, topic, last_checked, is_active."),
    ]
    tbl6 = doc.add_table(rows=1, cols=2)
    tbl6.style = "Table Grid"
    for cell, text in zip(tbl6.rows[0].cells, ["Table", "Purpose and Key Columns"]):
        cell.text = text
        for para in cell.paragraphs:
            for run in para.runs:
                run.bold = True
    for tname, desc in db_tables:
        row = tbl6.add_row()
        row.cells[0].text = tname
        row.cells[1].text = desc
    doc.add_paragraph()

    # ── 16. External API Integrations ──
    add_heading(doc, "16. External Service Integrations", level=1)

    ext_apis = [
        ("Google Gemini (via google-genai SDK)", "Primary AI backbone. Models: gemini-3.1-flash-lite-preview (utility), gemini-3.5-flash (chat/detector/research fallback), deep-research-preview-04-2026 (agentic research)."),
        ("APIMart Gateway", "Provides access to Claude Sonnet 4.5 and GPT-5.4 via a unified OpenAI-compatible API. Used exclusively in Research Mode's multi-model pipeline."),
        ("Tavily Search API", "Real-time web search. Used in: Web Search mode, Research Mode, Plagiarism Checker (multi-query), Deep Research Agent, Research Watcher."),
        ("Semantic Scholar API", "Free academic paper search API (200M+ papers). Used in the Plagiarism Checker to find academic source matches."),
        ("Firebase Authentication", "Handles user sign-up/login (Google OAuth and email+password). The Firebase UID is used as the bridge key to look up the Supabase user record."),
        ("Supabase", "PostgreSQL database and client library. Stores all user data, chat history, memories, documents, payments, and quota counters."),
        ("Razorpay", "Indian payment gateway. Handles order creation, checkout, HMAC verification, and webhooks for Plus/Pro subscriptions."),
        ("Brevo (formerly Sendinblue)", "Transactional email service. Used to send Research Watcher digest emails. The Brevo sending IP (35.234.215.209) is whitelisted in Supabase."),
        ("Pollinations AI", "Free image generation API. Used for the Image Generation feature (available to Plus/Pro users)."),
        ("Microsoft Edge TTS", "Text-to-speech synthesis via the edge-tts Python library. Used for the Radio Mode and voice output features."),
    ]
    tbl7 = doc.add_table(rows=1, cols=2)
    tbl7.style = "Table Grid"
    for cell, text in zip(tbl7.rows[0].cells, ["Service", "Role in Dynamo AI"]):
        cell.text = text
        for para in cell.paragraphs:
            for run in para.runs:
                run.bold = True
    for svc, desc in ext_apis:
        row = tbl7.add_row()
        row.cells[0].text = svc
        row.cells[1].text = desc
    doc.add_paragraph()

    # ── 17. Security Architecture ──
    add_heading(doc, "17. Security Architecture", level=1)
    security_points = [
        ("Backend plan enforcement", "All premium feature gates are enforced server-side. Frontend gates are UX conveniences only and cannot be relied upon for security."),
        ("HMAC-SHA256 payment verification", "Razorpay payment signatures are verified using a secret key before any plan upgrade is applied. This prevents forged payment confirmations."),
        ("No JWT exposure", "The Supabase service role key is only used server-side. The frontend uses the Supabase anonymous key for read-only user lookups."),
        ("Supabase Row Level Security", "Supabase tables use RLS policies to ensure users can only read/write their own records."),
        ("No server-side session storage", "The backend is stateless — all session context is passed in the request body (user_id, chat_id). The Deep Research Agent's in-memory job store is the only exception."),
        ("API key management", "All API keys (Gemini, Tavily, Razorpay, Brevo, Supabase, APIMart) are stored as environment variables and never exposed to the frontend."),
    ]
    for title_s, desc in security_points:
        p = doc.add_paragraph(style="List Bullet")
        p.add_run(title_s + ": ").bold = True
        p.add_run(desc).font.size = Pt(10)
    doc.add_paragraph()

    # ── 18. Conclusion ──
    add_heading(doc, "18. Conclusion", level=1)
    add_para(doc,
        "Dynamo AI demonstrates that a production-grade, multi-model AI platform can be built and "
        "operated by a small team with minimal infrastructure overhead. By serving both the frontend "
        "and backend from a single FastAPI process, integrating plan-aware model routing, and "
        "enforcing all business logic server-side, the system achieves both operational simplicity "
        "and commercial robustness.", size=10)
    add_para(doc,
        "The architectural decisions described in this paper — particularly the tiered model strategy, "
        "the two-layer memory filter, the multi-model research pipeline, and the in-house plagiarism "
        "detection system — represent replicable patterns for building cost-effective, feature-rich "
        "AI applications on top of commodity LLM APIs.", size=10)
    add_para(doc,
        "Future work includes the integration of Gemini 3.5 Pro (expected June 2026) for DeepThink "
        "upgrades, Gemini Omni Flash for video generation, and a planned React/Vite frontend migration "
        "for improved maintainability and component reuse.", size=10)

    doc.add_paragraph()

    # ── References ──
    add_heading(doc, "References", level=1)
    refs = [
        "[1] Google DeepMind. (2026). Gemini 3.5 Flash: Technical Report. Google AI Blog.",
        "[2] Anthropic. (2025). Claude Sonnet 4.5: Model Card. Anthropic Research.",
        "[3] OpenAI. (2025). GPT-5.4 System Card. OpenAI.",
        "[4] Supabase Inc. (2024). Supabase Documentation. https://supabase.com/docs",
        "[5] Google Firebase. (2024). Firebase Authentication Documentation. https://firebase.google.com/docs/auth",
        "[6] Razorpay. (2024). Payment Gateway API Documentation. https://razorpay.com/docs/api",
        "[7] Tavily AI. (2024). Tavily Search API Documentation. https://docs.tavily.com",
        "[8] Semantic Scholar. (2024). Semantic Scholar Open Research Corpus API. https://api.semanticscholar.org",
        "[9] Brevo. (2024). Transactional Email API. https://developers.brevo.com",
        "[10] FastAPI. (2024). FastAPI Documentation. https://fastapi.tiangolo.com",
    ]
    for ref in refs:
        p = doc.add_paragraph(style="List Bullet")
        p.add_run(ref).font.size = Pt(9)

    out_path = "/home/runner/workspace/Dynamo_AI_Architecture_Paper.docx"
    doc.save(out_path)
    print(f"Document saved: {out_path}")
    return out_path


if __name__ == "__main__":
    print("Generating architectural diagram...")
    diagram = make_diagram()
    print("Building Word document...")
    doc_path = build_doc(diagram)
    print(f"\nDone!\n  Diagram : {diagram}\n  Document: {doc_path}")
